import random
import numpy as np
import sys
import docx
from docx.shared import Mm

part = sys.argv[1]

# 配列に0が入っているか確認する関数
def zero_count(arraylist):
    for i in range(0, 9):
        for j in range(0, 9):
            n = num_81[(i, j)]
            if(n == 0):
                return True
            else:
                pass
    # 0が「num_81」に１つも入ってないときはFalseを返す
    return False

# 横軸に同じ数字が入っていないか確認する関数
def row_check(number, x, y):
    c = 0
    for i in range(0, 9):
        if(number == num_81[(x, i)]):
            c = c + 1
        else:
            pass
    if(c == 1):
        return True
    else:
        return False

# 縦軸に同じ数字が入っていないか確認する関数
def col_check(number, x, y):
    c = 0
    for i in range(0, 9):
        if(number == num_81[(i, y)]):
            c = c + 1
        else:
            pass
    if(c == 1):
        return True
    else:
        return False

# 3×3のボックスの中に同じ数字が入っていないことを確認する
def box_check(number, x, y):
    offset1 = [0, 1, 2]
    offset2 = [-1, 0, 1]
    offset3 = [-2, -1, 0]
    offset_list = [offset1, offset2, offset3]
    x_offset = offset_list[x % 3]
    y_offset = offset_list[y % 3]
    c = 0
    for i in range(0, 3):
        for j in range(0, 3):
            setx = x + x_offset[i]
            sety = y + y_offset[j]
            n = num_81[(setx, sety)]
            if(number == n):
                c = c + 1
            else:
                pass
    if(c == 1):
        return True
    else:
        return False

# 指定された座標から横軸に格納可能な数字のリストを返す
def row_num_list(x, y):
    num_list = [1, 2, 3, 4, 5, 6, 7, 8, 9]
    for i in range(0, 9):
        n = num_81[(x, i)]
        if( (y != i) and (n != 0) ):
            num_list.remove(n)
    return num_list

# 指定された座標から縦軸に格納可能な数字のリストを返す
def col_num_list(x, y):
    num_list = [1, 2, 3, 4, 5, 6, 7, 8, 9]
    for i in range(0, 9):
        n = num_81[(i, y)]
        if( (x != i) and (n != 0) ):
            num_list.remove(n)
    return num_list

# 指定された座標に対応する3*3boxの中に格納可能な数字のリストを返す
def box_num_list(x, y):
    num_list = [1, 2, 3, 4, 5, 6, 7, 8, 9]
    offset1 = [0, 1, 2]
    offset2 = [-1, 0, 1]
    offset3 = [-2, -1, 0]
    offset_list = [offset1, offset2, offset3]
    x_offset = offset_list[x % 3]
    y_offset = offset_list[y % 3]
    for i in range(0, 3):
        for j in range(0, 3):
            setx = x + x_offset[i]
            sety = y + y_offset[j]
            n = num_81[(setx, sety)]
            if( (x != setx or y != sety) and (n != 0) ):
                num_list.remove(n)
    return num_list

def ok_row_num(x, y):
    coordinate_list = [] # 座標がリストされている
    collect_list = []    # and_listをリストしている
    all_list = []        # and_listを全て足し合わせたもの
    for i in range(0, 9):
        # x, y = x, i
        # 空欄のマスに格納可能な数字のリストを得る
        if( num_81[(x, i)] == 0 ):
            row_list = row_num_list(x, i)
            col_list = col_num_list(x, i)
            box_list = box_num_list(x, i)
            l = set(row_list) & set(col_list) & set(box_list)
            and_list = list(l)
            coordinate_list.append((x, i))
            collect_list.append(and_list)
            all_list = all_list + and_list
    for i in range(0, len(coordinate_list)):
        for j in collect_list[i]:
            coorx, coory = coordinate_list[i]
            c = all_list.count(j)
            if( (x == coorx) and (y == coory) and (c == 1) ):
                return j
    return 0

def ok_col_num(x, y):
    coordinate_list = [] # 座標がリストされている
    collect_list = []    # and_listをリストしている
    all_list = []        # and_listを全て足し合わせたもの
    for i in range(0, 9):
        # x, y = i, y
        # 空欄のマスに格納可能な数字のリストを得る
        if( num_81[(i, y)] == 0 ):
            row_list = row_num_list(i, y)
            col_list = col_num_list(i, y)
            box_list = box_num_list(i, y)
            l = set(row_list) & set(col_list) & set(box_list)
            and_list = list(l)
            coordinate_list.append((i, y))
            collect_list.append(and_list)
            all_list = all_list + and_list
    for i in range(0, len(coordinate_list)):
        for j in collect_list[i]:
            coorx, coory = coordinate_list[i]
            c = all_list.count(j)
            if( (x == coorx) and (y == coory) and (c == 1) ):
                return j
    return 0

def ok_box_num(x, y):
    offset1 = [0, 1, 2]
    offset2 = [-1, 0, 1]
    offset3 = [-2, -1, 0]
    offset_list = [offset1, offset2, offset3]
    x_offset = offset_list[x % 3]
    y_offset = offset_list[y % 3]
    coordinate_list = []
    collect_list = []
    all_list = []
    for i in range(0, 3):
        for j in range(0, 3):
            # (x, y) = (setx, sety)
            setx = x + x_offset[i]
            sety = y + y_offset[j]
            if( num_81[(setx, sety)] == 0 ):
                row_list = row_num_list(setx, sety)
                col_list = col_num_list(setx, sety)
                box_list = box_num_list(setx, sety)
                l = set(row_list) & set(col_list) & set(box_list)
                and_list = list(l)
                coordinate_list.append((setx, sety))
                collect_list.append(and_list)
                all_list = all_list + and_list
    # リストから数字が1回しか出現しないリストを特定する
    # そして格納する
    for i in range(0, len(coordinate_list)):
        for j in collect_list[i]:
            coorx, coory = coordinate_list[i]
            c = all_list.count(j)
            if( (x == coorx) and (y == coory) and (c == 1) ):
                return j
    return 0

# 上記3つの関数から得られるリストの共通部を出力し
# 新たにリストを作成。　リストの中からランダムに数字を1つ返す
def ok_num(x, y):
    row_list = row_num_list(x, y)
    col_list = col_num_list(x, y)
    box_list = box_num_list(x, y)
    l = set(row_list) & set(col_list) & set(box_list)
    and_list = list(l)
    if( len(and_list) == 0):
        return 0
    else:
        row_num = ok_row_num(x, y)
        col_num = ok_col_num(x, y)
        box_num = ok_box_num(x, y)
        row_num_TF = row_num in and_list
        col_num_TF = col_num in and_list
        box_num_TF = box_num in and_list
        if( row_num_TF == True ):
            return row_num
        else:
            if(col_num_TF == True):
                return col_num
            else:
                if(box_num_TF == True):
                    return box_num
                else:
                    r = random.randrange(0, len(and_list))
                    return and_list[r]


# 横軸において,指定された数字が他のマスで唯一の数字になっていないか確認する
def row_only_check(number, x, y):
    row_only_list = []
    for i in range(0, 9):
        row_list = row_num_list(x, i)
        col_list = col_num_list(x, i)
        box_list = box_num_list(x, i)
        l = set(row_list) & set(col_list) & set(box_list)
        and_list = list(l)
        row_only_list.append(and_list)
    for i in range(0, 9):
        rolist = row_only_list[i]
        if( (y != i) and (len(rolist) == 1) and (rolist[0] == number) ):
            return False
    return True

#　縦軸において,指定された数字が他のマスで唯一の数字になっていないか確認する
def col_only_check(number, x, y):
    col_only_list = []
    for i in range(0, 9):
        row_list = row_num_list(i, y)
        col_list = col_num_list(i, y)
        box_list = box_num_list(i, y)
        l = set(row_list) & set(col_list) & set(box_list)
        and_list = list(l)
        col_only_list.append(and_list)
    for i in range(0, 9):
        colist = col_only_list[i]
        if( (x != i) and (len(colist) == 1) and (colist[0] == number) ):
            return False
    return True

# 3*3のボックス内において,指定された数字が他のマスで唯一の数字になっていないか確認する
def box_only_check(number, x, y):
    box_only_list = []
    offset1 = [0, 1, 2]
    offset2 = [-1, 0, 1]
    offset3 = [-2, -1, 0]
    offset_list = [offset1, offset2, offset3]
    x_offset = offset_list[x % 3]
    y_offset = offset_list[y % 3]
    for i in range(0, 3):
        for j in range(0, 3):
            setx = x + x_offset[i]
            sety = y + y_offset[j]
            row_list = row_num_list(setx, sety)
            col_list = col_num_list(setx, sety)
            box_list = box_num_list(setx, sety)
            l = set(row_list) & set(col_list) & set(box_list)
            and_list = list(l)
            box_only_list.append(and_list)
    count = 0
    for i in range(0, 3):
        for j in range(0, 3):
            bolist = box_only_list[count]
            count = count + 1
            setx = x + x_offset[i]
            sety = y + y_offset[j]
            if( (x != setx or y != sety) and (len(bolist) == 1) and (bolist[0] == number) ):
                return False
    return True

# 左上の3*3マスと真ん中の3*3マスと右下の3*3マスに
# ランダムに数字を代入する
# place = 0 or 3 or 6
def rand_insert(place, emp, num_81_list):
    num_list = [1, 2, 3, 4, 5, 6, 7, 8, 9]
    random.shuffle(num_list)
    emp_li = []
    for i in range(0, 3):
        for j in range(0, 3):
            emp_li.append((i + place, j + place))
    while(len(emp_li) > 0):
        r = random.randrange(0, len(emp_li))
        x, y = emp_li[r]
        count = 0
        while(True):
            try:
                empx, empy = emp[count]
            except IndexError:
                print("rand_insert:count = ",count)
                print("rand_insert:len(emp) = ", len(emp))
                sys.exit()
            if( (x == empx) and (y == empy) ):
                num_81_list[(x, y)] = num_list[r]
                erasure1 = emp.pop(count)
                erasure2 = emp_li.pop(r)
                number = num_list.pop(r)
                break
            else:
                count = count + 1
    return emp

# check_list True or False
def check_list(n, x, y):
    rc_TF = row_check(n, x, y)
    cc_TF = col_check(n, x, y)
    bc_TF = box_check(n, x, y)
    roc_TF = row_only_check(n, x, y);
    coc_TF = col_only_check(n, x, y);
    boc_TF = box_only_check(n, x, y);
    if( (rc_TF == True) and (cc_TF == True) and (bc_TF == True) and (roc_TF == True) and (coc_TF == True) and (boc_TF == True)):
        return True
    return False

# 横に置いて各マスに格納可能な数字のリスト作成する
# 既に格納されている場合は空値を返す
# リストを足して数字が1回しか出てこなかったらその数字を代入する
def row_only_set(emp, num_81):
    for i in range(0, 9):
        coordinate_list = []
        collect_list = []
        all_list = []
        for j in range(0, 9):
            # x, y = i, j
            # 空欄のマスに格納可能な数字のリストを得る
            if( num_81[(i, j)] == 0 ):
                row_list = row_num_list(i, j)
                col_list = col_num_list(i, j)
                box_list = box_num_list(i, j)
                l = set(row_list) & set(col_list) & set(box_list)
                and_list = list(l)
                coordinate_list.append((i, j))
                collect_list.append(and_list)
                all_list = all_list + and_list
        # リストから数字が1回しか出現しないリストを特定する
        # そして格納する
        for k in range(0, len(all_list)):
            c = all_list.count(all_list[k])
            if( c == 1 ):
                for coor in range(0, len(collect_list)):
                    TF = all_list[k] in collect_list[coor]
                    if( TF == True ):
                        coorx, coory = coordinate_list[coor]
                        num_81[(coorx, coory)] = all_list[k]
                        count_up = 0
                        clist = len(emp)
                        while( clist > count_up ):
                            empx, empy = emp[count_up]
                            if( (coorx == empx) and (coory == empy) ):
                                erasure = emp.pop(count_up)
                                count_up = count_up - 1
                                break
                            else:
                                pass
                            count_up = count_up + 1
                            clist = len(emp)
    return emp

# 縦に置いて各マスに格納可能な数字のリスト作成する
# 既に格納されている場合は空値を返す
# リストを足して数字が1回しか出てこなかったらその数字を代入する
def col_only_set(emp, num_81):
    for i in range(0, 9):
        coordinate_list = []
        collect_list = []
        all_list = []
        for j in range(0, 9):
            # x, y = j, i
            # 空欄のマスに格納可能な数字のリストを得る
            if( num_81[(j, i)] == 0 ):
                row_list = row_num_list(j, i)
                col_list = col_num_list(j, i)
                box_list = box_num_list(j, i)
                l = set(row_list) & set(col_list) & set(box_list)
                and_list = list(l)
                coordinate_list.append((j, i))
                collect_list.append(and_list)
                all_list = all_list + and_list
        # リストから数字が1回しか出現しないリストを特定する
        # そして格納する
        for k in range(0, len(all_list)):
            c = all_list.count(all_list[k])
            if( c == 1 ):
                for coor in range(0, len(collect_list)):
                    TF = all_list[k] in collect_list[coor]
                    if( TF == True ):
                        coorx, coory = coordinate_list[coor]
                        num_81[(coorx, coory)] = all_list[k]
                        count_up = 0
                        clist = len(emp)
                        while( clist > count_up ):
                            empx, empy = emp[count_up]
                            if( (coorx == empx) and (coory == empy) ):
                                erasure = emp.pop(count_up)
                                count_up = count_up - 1
                                break
                            else:
                                pass
                            count_up = count_up + 1
                            clist = len(emp)
    return emp

# 3*3boxに置いて各マスに格納可能な数字のリスト作成する
# 既に格納されている場合は空値を返す
# リストを足して数字が1回しか出てこなかったらその数字を代入する
def box_only_set(emp, num_81):
    offset1 = [0, 1, 2]
    offset2 = [-1, 0, 1]
    offset3 = [-2, -1, 0]
    offset_list = [offset1, offset2, offset3]
    count_all = 0
    clist = len(emp)
    while( clist > count_all ):
        boxx, boxy = emp[count_all]
        x_offset = offset_list[boxx % 3]
        y_offset = offset_list[boxy % 3]
        coordinate_list = []
        collect_list = []
        all_list = []
        for i in range(0, 3):
            for j in range(0, 3):
                # (x, y) = (setx, sety)
                setx = boxx + x_offset[i]
                sety = boxy + y_offset[j]
                if( num_81[(setx, sety)] == 0 ):
                    row_list = row_num_list(setx, sety)
                    col_list = col_num_list(setx, sety)
                    box_list = box_num_list(setx, sety)
                    l = set(row_list) & set(col_list) & set(box_list)
                    and_list = list(l)
                    coordinate_list.append((setx, sety))
                    collect_list.append(and_list)
                    all_list = all_list + and_list
        # リストから数字が1回しか出現しないリストを特定する
        # そして格納する
        for k in range(0, len(all_list)):
            c = all_list.count(all_list[k])
            if( c == 1 ):
                for coor in range(0, len(collect_list)):
                    TF = all_list[k] in collect_list[coor]
                    if( TF == True ):
                        coorx, coory = coordinate_list[coor]
                        num_81[(coorx, coory)] = all_list[k]
                        count_up = 0
                        clist = len(emp)
                        while( clist > count_up ):
                            empx, empy = emp[count_up]
                            if( (coorx == empx) and (coory == empy) ):
                                erasure = emp.pop(count_up)
                                count_up = count_up - 1
                                break
                            else:
                                pass
                            count_up = count_up + 1
                            clist = len(emp)
        count_all = count_all + 1
        clist = len(emp)
    return emp

# 上記のrow and col and boxを統合する関数
def all_num_set(emp, num_81):
    while(True):
        row_emp = row_only_set(emp, num_81)
        col_emp = col_only_set(emp, num_81)
        box_emp = box_only_set(emp, num_81)
        if( row_emp == col_emp == box_emp ):
            break
    return emp

# 9*9の表を表示する,9*9の配列を引数に入れてください
def table_print(list):
    print(" +----------+----------+----------+")
    for i in range(0, 3):
        for j in range(0, 3):
            print(" | ", end="")
            for k in range(0, 3):
                print(int(list[(j+3*i, k)]), " ",  end="")
            print("| ", end="")
            for k in range(3, 6):
                print(int(list[(j+3*i, k)]), " ",  end="")
            print("| ", end="")
            for k in range(6, 9):
                print(int(list[(j+3*i, k)]), " ", end="")
            print("|")
        print(" +----------+----------+----------+")


# 変数定義 ======================================#
# 9行9列の配列を生成
num_81 = np.zeros((9,9), dtype=int)
# 空白の座標リスト（正しくは0が格納されている座標リスト）
emp = []
for i in range(0, 9):
    for j in range(0, 9):
        emp.append((i, j))
#===============================================#

# ここからmainの実行部分
while(len(emp) > 0):
    c = 0
    emp = rand_insert(0, emp, num_81)
    emp = rand_insert(3, emp, num_81)
    emp = rand_insert(6, emp, num_81)
    while(c < 81):
        c = c + 1
        zc_TF = zero_count(num_81)
        emp_l = len(emp)
        if( (zc_TF == True) and (emp_l >= 0) ):
            emp = all_num_set(emp, num_81)
            emp_l = len(emp)
            if(emp_l == 0):
                break
            r = random.randint(0, emp_l - 1)
            x, y = emp[r]
            n = ok_num(x, y)
            if(n != 0):
                num_81[emp[r]] = n
                cl_TF = check_list(n, x, y)
                if( (cl_TF == True) ):
                    emp.pop(r)
                else:
                    num_81[emp[r]] = 0
                    pass
            else:
                break
        else:
            break
    if(len(emp) == 0):
        break
    # 再定義
    num_81 = np.zeros((9,9), dtype=int)
    emp = []
    for i in range(0, 9):
        for j in range(0, 9):
            emp.append((i, j))

# ここからHTMLに書き込む部分
html = """\
<!DOCTYPE html>
<html lang="ja" dir="ltr">
  <head>
    <meta charset="utf-8">
    <title>数独</title>
    <link rel="stylesheet" href="style.css" media="screen">
    <link rel="stylesheet" href="print.css" media="print">

  </head>
  <body>
    <div class="main-body">
      {0}
      <br>
    </div>
  </body>
</html>
"""

answer_parts = """\
<div class="answer">
  <h3>答え</h3>
  <div class="row">
    <!-- 1段目 -->
    {0}
    {1}
    {2}
  </div>
  <div class="row">
    <!-- 2段目 -->
    {3}
    {4}
    {5}
  </div>
  <div class="row">
    <!-- 3段目 -->
    {6}
    {7}
    {8}
  </div>
</div>
"""

table_parts = """
        <table class="flex-table">
          <tr><th>{0}</th><th>{1}</th><th>{2}</th></tr>
          <tr><th>{3}</th><th>{4}</th><th>{5}</th></tr>
          <tr><th>{6}</th><th>{7}</th><th>{8}</th></tr>
        </table>
"""

answer_list = []

# ここから　マスを空欄にするための部分==============================
index_list = [i for i in range(0, 81)]
xy_list = []
emp = []
for i in range(0, 9):
    for j in range(0, 9):
        emp.append((i, j))

count = 0
while(len(index_list) > 0):
    r = random.randrange(0, len(index_list))
    xy_list.append(emp[index_list[r]])
    index_list.pop(r)
# ここまで ======================================================

for i in range(0, 9, 3):
    for j in range(0, 9, 3):
        al = table_parts.format(num_81[(i, j)], num_81[(i, j+1)], num_81[(i, j+2)], num_81[(i+1, j)], num_81[(i+1, j+1)], num_81[(i+1, j+2)], num_81[(i+2, j)], num_81[(i+2, j+1)], num_81[(i+2, j+2)])
        answer_list.append(al)

answer_parts = answer_parts.format(answer_list[0], answer_list[1], answer_list[2], answer_list[3], answer_list[4], answer_list[5], answer_list[6], answer_list[7], answer_list[8])
answer = html.format(answer_parts)

answer_path = f"answer_part{part}.html"
with open(answer_path, mode="w", encoding="utf-8") as f:
    f.write(answer)

#ここからヒント用紙を製作する部分
doc = docx.Document()
doc.sections[0].left_margin = Mm(10)
doc.sections[0].right_margin= Mm(10)
doc.sections[0].top_margin= Mm(6)
doc.sections[0].bottom_margin= Mm(6)
for i in range(9):
    i12=i*13
    doc.add_paragraph('　')
    doc.add_paragraph(f'{i}行0列目は      {i}行1列目は      {i}行2列目は')
    doc.add_paragraph('　')
    doc.add_paragraph(f'   {num_81[i][0]}         {num_81[i][1]}         {num_81[i][2]}  ')
    doc.paragraphs[i12+0].runs[0].font.size = docx.shared.Pt(20)
    doc.paragraphs[i12+1].runs[0].font.size = docx.shared.Pt(30)
    doc.paragraphs[i12+2].runs[0].font.size = docx.shared.Pt(12)
    doc.paragraphs[i12+3].runs[0].font.size = docx.shared.Pt(80)
    doc.add_paragraph('　')
    doc.add_paragraph(f'{i}行3列目は      {i}行4列目は      {i}行5列目は')
    doc.add_paragraph('　')
    doc.add_paragraph(f'   {num_81[i][3]}         {num_81[i][4]}         {num_81[i][5]}  ')
    doc.paragraphs[i12+4].runs[0].font.size = docx.shared.Pt(20)
    doc.paragraphs[i12+5].runs[0].font.size = docx.shared.Pt(30)
    doc.paragraphs[i12+6].runs[0].font.size = docx.shared.Pt(12)
    doc.paragraphs[i12+7].runs[0].font.size = docx.shared.Pt(80)
    doc.add_paragraph('　')
    doc.add_paragraph(f'{i}行6列目は      {i}行7列目は      {i}行8列目は')
    doc.add_paragraph('　')
    doc.add_paragraph(f'   {num_81[i][6]}         {num_81[i][7]}         {num_81[i][8]}  ')
    doc.paragraphs[i12+8].runs[0].font.size = docx.shared.Pt(20)
    doc.paragraphs[i12+9].runs[0].font.size = docx.shared.Pt(30)
    doc.paragraphs[i12+10].runs[0].font.size = docx.shared.Pt(12)
    doc.paragraphs[i12+11].runs[0].font.size = docx.shared.Pt(80)
    if i!=8:
        doc.add_page_break()
doc.save(f"hint_part{part}.docx")